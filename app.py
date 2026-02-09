import os
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Skip model source connectivity check and set GPU device
os.environ['DISABLE_MODEL_SOURCE_CHECK'] = 'True'
os.environ['CUDA_VISIBLE_DEVICES'] = '0'

import tempfile
import uuid
import json
import threading
import time
import shutil
import atexit
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory, send_file, Response, after_this_request
from flask_wtf.csrf import CSRFProtect, generate_csrf
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import fitz  # PyMuPDF
import numpy as np

# Set Paddle device early
import paddle
paddle.set_device('gpu:0')

# Configuration constants (can be overridden via environment)
MAX_UPLOAD_SIZE = int(os.environ.get('PDF_OCR_MAX_UPLOAD_MB', 500)) * 1024 * 1024
CLEANUP_INTERVAL = int(os.environ.get('PDF_OCR_CLEANUP_INTERVAL', 3600))
MAX_FILE_AGE = int(os.environ.get('PDF_OCR_MAX_FILE_AGE', 3600))
MODEL_IDLE_TIMEOUT = int(os.environ.get('PDF_OCR_MODEL_IDLE_TIMEOUT', 1800))
JOB_MAX_AGE = int(os.environ.get('PDF_OCR_JOB_MAX_AGE', 7200))
OCR_BATCH_SIZE = 4
PANDOC_TIMEOUT = 120
PDF_POINTS_PER_INCH = 72

app = Flask(__name__, static_folder='static')
app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_SIZE
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(32).hex())
app.config['WTF_CSRF_TIME_LIMIT'] = None  # No expiry for SPA

# CSRF protection
csrf = CSRFProtect(app)

# Rate limiting
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["100 per minute"],
    storage_uri="memory://"
)

# Store processed files temporarily
OUTPUT_DIR = Path(tempfile.gettempdir()) / 'pdf_ocr_output'
OUTPUT_DIR.mkdir(exist_ok=True)

# Upload directory for async processing
UPLOAD_DIR = Path(tempfile.gettempdir()) / 'pdf_ocr_uploads'
UPLOAD_DIR.mkdir(exist_ok=True)

# Job tracking
jobs = {}
jobs_lock = threading.Lock()

# Progress tracking (kept for SSE compatibility)
progress_data = {}
progress_lock = threading.Lock()

# Cancel flags for jobs
cancel_flags = {}
cancel_lock = threading.Lock()

# Processing locks - separate locks for OCR vs VL pipelines to allow some parallelism
# Note: PaddlePaddle models are not thread-safe, but OCR and VL use different models
ocr_processing_lock = threading.Lock()  # For searchable PDF (OCR pipeline)
vl_processing_lock = threading.Lock()   # For Markdown/Word (VL pipeline)

# Cleanup state
cleanup_thread = None
cleanup_stop_event = threading.Event()

# Initialize PaddleOCR (lazy loading)
_ocr = None
_vl_pipeline = None
_model_lock = threading.Lock()

# Model idle tracking for VRAM release
_ocr_last_used = 0
_vl_last_used = 0


def cleanup_old_files():
    """Background thread to clean up old exported files."""
    while not cleanup_stop_event.is_set():
        try:
            now = time.time()
            # Clean up old PDF outputs
            for f in OUTPUT_DIR.glob('*.pdf'):
                if f.is_file() and (now - f.stat().st_mtime) > MAX_FILE_AGE:
                    try:
                        f.unlink()
                        logger.info(f"Cleanup: Deleted old file: {f.name}")
                    except Exception as e:
                        logger.info(f"Cleanup: Error deleting {f.name}: {e}")

            # Clean up old folders (markdown/word outputs)
            for d in OUTPUT_DIR.iterdir():
                if d.is_dir() and (now - d.stat().st_mtime) > MAX_FILE_AGE:
                    try:
                        shutil.rmtree(d)
                        logger.info(f"Cleanup: Deleted old folder: {d.name}")
                    except Exception as e:
                        logger.info(f"Cleanup: Error deleting {d.name}: {e}")

            # Clean up old zip files
            for f in OUTPUT_DIR.glob('*.zip'):
                if f.is_file() and (now - f.stat().st_mtime) > MAX_FILE_AGE:
                    try:
                        f.unlink()
                        logger.info(f"Cleanup: Deleted old zip: {f.name}")
                    except Exception as e:
                        logger.info(f"Cleanup: Error deleting {f.name}: {e}")

            # Release idle models from VRAM (check every cleanup cycle)
            try:
                release_idle_models()
            except Exception as e:
                logger.info(f"Cleanup: Error releasing idle models: {e}")

            # Clean up old jobs from memory (older than JOB_MAX_AGE)
            try:
                with jobs_lock:
                    stale_jobs = [jid for jid, job in jobs.items()
                                  if now - job.get('updated_at', job.get('created_at', 0)) > JOB_MAX_AGE]
                    for jid in stale_jobs:
                        del jobs[jid]
                        logger.info(f"Cleanup: Removed stale job: {jid}")

                with progress_lock:
                    # Clean up progress data for stale jobs
                    stale_progress = [tid for tid in progress_data if tid in stale_jobs]
                    for tid in stale_progress:
                        del progress_data[tid]
            except Exception as e:
                logger.info(f"Cleanup: Error cleaning up jobs: {e}")

        except Exception as e:
            logger.info(f"Cleanup: Error during cleanup: {e}")

        # Wait for next cleanup cycle
        cleanup_stop_event.wait(CLEANUP_INTERVAL)


def start_cleanup_thread():
    """Start the background cleanup thread."""
    global cleanup_thread
    if cleanup_thread is None or not cleanup_thread.is_alive():
        cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
        cleanup_thread.start()
        logger.info("Cleanup: Background cleanup thread started")


def stop_cleanup_thread():
    """Stop the background cleanup thread."""
    cleanup_stop_event.set()
    if cleanup_thread and cleanup_thread.is_alive():
        cleanup_thread.join(timeout=5)
        logger.info("Cleanup: Background cleanup thread stopped")


# Register cleanup on exit
atexit.register(stop_cleanup_thread)

# Start cleanup thread
start_cleanup_thread()


def schedule_file_deletion(path: Path, delay: float = 5.0):
    """Schedule a file or directory for deletion after a delay."""
    def delete_later():
        time.sleep(delay)
        try:
            if path.is_dir():
                shutil.rmtree(path)
                logger.info(f"Cleanup: Deleted after download: {path.name}/")
            elif path.exists():
                path.unlink()
                logger.info(f"Cleanup: Deleted after download: {path.name}")
        except Exception as e:
            logger.info(f"Cleanup: Error deleting {path}: {e}")

    thread = threading.Thread(target=delete_later, daemon=True)
    thread.start()


def get_ocr():
    global _ocr, _ocr_last_used
    with _model_lock:
        if _ocr is None:
            from paddleocr import PaddleOCR
            _ocr = PaddleOCR(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                enable_hpi=True,  # ONNX Runtime GPU acceleration
                # PP-OCRv5 server models (5-10% better accuracy)
                text_detection_model_name="PP-OCRv5_server_det",
                text_recognition_model_name="PP-OCRv5_server_rec",
                # Batch size optimized for 8GB+ VRAM
                text_recognition_batch_size=16,
                # Detection tuning for better accuracy
                text_det_limit_side_len=960,
                text_det_thresh=0.25,
                text_det_box_thresh=0.5,
                # CPU optimization for preprocessing
                cpu_threads=8,
                enable_mkldnn=True,
            )
        _ocr_last_used = time.time()
    return _ocr


def get_vl_pipeline():
    """Get PaddleOCR-VL-1.5 pipeline for document parsing."""
    global _vl_pipeline, _vl_last_used
    with _model_lock:
        if _vl_pipeline is None:
            from paddleocr import PaddleOCRVL
            _vl_pipeline = PaddleOCRVL(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                # Note: enable_hpi not supported for VL pipeline's layout models
                cpu_threads=8,
                enable_mkldnn=True,
            )
            logger.info("VL: Initialized VL pipeline")
        _vl_last_used = time.time()
    return _vl_pipeline


def release_idle_models():
    """Release models from VRAM if unused for MODEL_IDLE_TIMEOUT seconds.

    Called periodically by the cleanup thread to free GPU memory when models
    are not actively being used. Frees 2-4GB VRAM when models are released.
    """
    global _ocr, _vl_pipeline
    current_time = time.time()

    with _model_lock:
        if _ocr is not None and (current_time - _ocr_last_used) > MODEL_IDLE_TIMEOUT:
            del _ocr
            _ocr = None
            import gc
            gc.collect()
            try:
                paddle.device.cuda.empty_cache()
            except Exception:
                pass
            logger.info("Memory: OCR model released from VRAM after idle timeout")

        if _vl_pipeline is not None and (current_time - _vl_last_used) > MODEL_IDLE_TIMEOUT:
            del _vl_pipeline
            _vl_pipeline = None
            import gc
            gc.collect()
            try:
                paddle.device.cuda.empty_cache()
            except Exception:
                pass
            logger.info("Memory: VL model released from VRAM after idle timeout")


# 簡體轉繁體修正（只轉換一對一的字，略過一對多）
from s2t_dict import S2T_ONE_TO_ONE

# Pre-build translation table for O(1) character conversion
_S2T_TABLE = str.maketrans(S2T_ONE_TO_ONE)

# Pre-compiled regex patterns for markdown processing (avoid recompilation per call)
import re as _re
_TABLE_PATTERNS = [
    _re.compile(r'<div[^>]*>\s*<html>\s*<body>\s*<table[^>]*>.*?</table>\s*</body>\s*</html>\s*</div>', _re.DOTALL),
    _re.compile(r'<div[^>]*class="[^"]*table[^"]*"[^>]*>.*?</div>', _re.DOTALL),
    _re.compile(r'<table[^>]*>.*?</table>', _re.DOTALL),
]
_IMG_PATTERNS = [
    _re.compile(r'<div[^>]*>\s*<img[^>]+/?>\s*</div>', _re.DOTALL),
    _re.compile(r'<img[^>]+/?>', _re.DOTALL),
]


def fix_ocr_text(text: str) -> str:
    """修正 OCR 誤認的簡體字，只轉換一對一的字，略過一對多。"""
    return text.translate(_S2T_TABLE)


def validate_pdf_file(file) -> bool:
    """Validate file is actually a PDF by checking magic bytes."""
    file.seek(0)
    header = file.read(5)
    file.seek(0)
    return header == b'%PDF-'


def validate_path(filename: str) -> Path:
    """Validate that filename doesn't escape OUTPUT_DIR (path traversal prevention)."""
    # Resolve the full path
    resolved = (OUTPUT_DIR / filename).resolve()
    # Ensure it's within OUTPUT_DIR using is_relative_to() for robust validation
    if not resolved.is_relative_to(OUTPUT_DIR.resolve()):
        raise ValueError(f"Invalid path: {filename}")
    return resolved


def sanitize_download_name(filename: str, max_length: int = 200) -> str:
    """Sanitize filename for use in HTTP Content-Disposition header.

    Removes or replaces characters that could cause issues:
    - Control characters (including newlines, tabs)
    - Quotes (could break header parsing)
    - Path separators (security)
    - Non-ASCII characters are preserved but could be encoded by Flask
    - Truncates to max_length to prevent HTTP header issues

    Args:
        filename: The original filename
        max_length: Maximum length for the filename (default 200)

    Returns:
        Sanitized filename safe for HTTP headers
    """
    import re
    # Remove control characters (0x00-0x1F, 0x7F)
    filename = re.sub(r'[\x00-\x1f\x7f]', '', filename)
    # Replace characters problematic in Content-Disposition
    # Quotes break the header, path separators are security risks
    filename = filename.replace('"', "'").replace('\\', '_').replace('/', '_')
    # Collapse multiple underscores/spaces
    filename = re.sub(r'[_\s]+', '_', filename)
    # Remove leading/trailing whitespace and dots (hidden files, extension tricks)
    filename = filename.strip(' ._')
    # Truncate to max length to prevent HTTP header issues
    if len(filename) > max_length:
        filename = filename[:max_length].rstrip(' ._')
    # Ensure we have something left
    if not filename:
        filename = 'download'
    return filename


def process_pdf_with_vl(input_pdf_path: str, output_dir: Path, task_id: str,
                        original_filename: str = None,
                        extra_steps: int = 0) -> tuple[list, Path, int, str]:
    """
    Shared VL processing logic for both Markdown and Word export.
    Returns (restructured_results, file_output_dir, total_pages, download_id).

    The download_id includes task_id to prevent directory collisions when
    multiple files with the same name are uploaded concurrently.

    Args:
        extra_steps: Additional steps after VL processing (e.g., 2 for Word conversion).
                     Used to calculate consistent progress percentage throughout the pipeline.
    """
    import fitz

    # Open and validate PDF (same checks as OCR pipeline)
    try:
        doc = fitz.open(input_pdf_path)
    except Exception as e:
        raise RuntimeError(f"無法開啟 PDF 檔案: {e}")

    # Check for encryption
    if doc.is_encrypted:
        doc.close()
        raise RuntimeError("無法處理加密的 PDF 檔案")

    total_pages = len(doc)
    if total_pages == 0:
        doc.close()
        raise RuntimeError("PDF 檔案沒有頁面")

    # Use total_steps for consistent progress tracking throughout the pipeline
    total_steps = total_pages + extra_steps
    update_progress(task_id, 0, total_steps, 'processing', f'開始轉換 {total_pages} 頁...')

    pipeline = get_vl_pipeline()

    # Create output directory for this file
    # Include task_id to prevent collisions when same filename is uploaded multiple times
    if original_filename:
        base_name = Path(original_filename).stem
    else:
        base_name = Path(input_pdf_path).stem
    download_id = f"{base_name}_{task_id}"
    file_output_dir = output_dir / download_id
    file_output_dir.mkdir(parents=True, exist_ok=True)

    doc_closed = False
    try:
        # Process each page individually for progress tracking
        pages_res = []
        for page_num in range(total_pages):
            # Check for cancellation
            if is_cancelled(task_id):
                raise RuntimeError("已取消處理")

            update_progress(task_id, page_num + 1, total_steps, 'processing',
                           f'辨識第 {page_num + 1}/{total_pages} 頁...')

            # Render page to image
            # 1.5x zoom: 平衡品質與記憶體 (VL 模型會自行處理解析度)
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))

            # Convert pixmap to numpy array directly (no temp file I/O)
            img_array = pixmap_to_numpy(pix)

            # Process single page image - VL pipeline accepts numpy arrays
            # 準確度優先：不限制 max_pixels
            output = pipeline.predict(input=img_array)
            for res in output:
                pages_res.append(res)

            # Free memory
            del img_array
            del pix

        # Close doc after processing all pages
        doc.close()
        doc_closed = True

        update_progress(task_id, total_pages, total_steps, 'saving', '整合頁面內容...')

        # Restructure pages with cross-page table merging and title releveling
        restructured = pipeline.restructure_pages(
            pages_res,
            merge_tables=True,
            relevel_titles=True,
            concatenate_pages=True
        )

        return restructured, file_output_dir, total_pages, download_id

    except Exception:
        # Only close doc if not already closed
        if not doc_closed:
            doc.close()
        # Clean up output directory on failure
        if file_output_dir.exists():
            shutil.rmtree(file_output_dir)
        raise


def update_progress(task_id: str, current: int, total: int, status: str, message: str = ""):
    with progress_lock:
        progress_data[task_id] = {
            'current': current,
            'total': total,
            'percent': min(100, int((current / total) * 100)) if total > 0 else 0,
            'status': status,
            'message': message,
            'updated_at': time.time()  # Timestamp for stale detection
        }


def update_job(job_id: str, **kwargs):
    """Update job status."""
    with jobs_lock:
        if job_id in jobs:
            jobs[job_id].update(kwargs)
            jobs[job_id]['updated_at'] = time.time()


def is_cancelled(job_id: str) -> bool:
    """Check if a job has been cancelled."""
    with cancel_lock:
        return cancel_flags.get(job_id, False)


def set_cancelled(job_id: str, value: bool = True):
    """Set cancel flag for a job."""
    with cancel_lock:
        cancel_flags[job_id] = value


def pixmap_to_numpy(pix: fitz.Pixmap) -> np.ndarray:
    """Convert PyMuPDF Pixmap to numpy array (RGB format for PaddleOCR).

    Args:
        pix: PyMuPDF Pixmap object

    Returns:
        numpy array in RGB format (H, W, 3)
    """
    # Get image data as bytes
    if pix.alpha:
        # RGBA -> RGB: drop alpha channel
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 4)
        img = img[:, :, :3]  # Drop alpha
    else:
        # Already RGB
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)

    # PaddleOCR expects BGR format (like OpenCV), but samples are RGB
    # Convert RGB to BGR
    img = img[:, :, ::-1].copy()

    return img


def parse_ocr_result(res: dict, min_confidence: float) -> list:
    """Parse single OCR result dict into standardized format.

    Args:
        res: Single result dict from PaddleOCR predict()
        min_confidence: Minimum confidence score to include a text block

    Returns:
        List of OCR data dicts with text, poly, and score
    """
    ocr_data = []
    if 'rec_texts' not in res or not res['rec_texts']:
        return ocr_data

    texts = res['rec_texts']
    polys = res.get('rec_polys', res.get('dt_polys', []))
    scores = res.get('rec_scores', [])

    for i, text in enumerate(texts):
        if i < len(polys):
            poly = polys[i]
            score = scores[i] if i < len(scores) else 1.0

            # Filter out low-confidence results
            if score < min_confidence:
                continue

            # Skip empty or whitespace-only text
            if not text or not text.strip():
                continue

            ocr_data.append({
                'text': fix_ocr_text(text),
                'poly': poly.tolist() if hasattr(poly, 'tolist') else poly,
                'score': score
            })

    return ocr_data


def ocr_image(image, min_confidence: float = 0.5) -> list:
    """Run OCR on an image and return text with bounding boxes.

    Args:
        image: Either a file path (str) or numpy array (H, W, 3) in BGR format
        min_confidence: Minimum confidence score to include a text block (0.0-1.0)

    Returns:
        List of OCR results with text, polygon coordinates, and confidence scores
    """
    ocr = get_ocr()

    try:
        result = ocr.predict(image)
    except Exception as e:
        logger.warning(f"OCR: OCR failed on image: {e}")
        return []

    ocr_data = []
    for res in result:
        ocr_data.extend(parse_ocr_result(res, min_confidence))

    return ocr_data


def ocr_batch_images(images: list, min_confidence: float = 0.5) -> list:
    """Run OCR on a batch of images for improved throughput.

    Processes multiple images in a single OCR call, leveraging GPU parallelism
    for 30-50% faster processing on multi-page documents.

    Args:
        images: List of numpy arrays (H, W, 3) in BGR format
        min_confidence: Minimum confidence score to include a text block (0.0-1.0)

    Returns:
        List of OCR results per image (same order as input)
    """
    if not images:
        return []

    ocr = get_ocr()

    try:
        # PaddleOCR predict() accepts a list of images for batch processing
        results = ocr.predict(images)
    except Exception as e:
        logger.warning(f"OCR: Batch OCR failed: {e}")
        # Fallback to individual processing
        return [ocr_image(img, min_confidence) for img in images]

    # Results come back as a flat list, one result dict per image
    batch_ocr_data = []
    for res in results:
        batch_ocr_data.append(parse_ocr_result(res, min_confidence))

    return batch_ocr_data


def create_searchable_pdf(input_pdf_path: str, output_pdf_path: str, task_id: str,
                          dpi: int = 200, min_confidence: float = 0.5) -> dict:
    """Create a searchable PDF by adding invisible text layer to original PDF.

    Uses batch OCR processing for 30-50% faster multi-page documents.

    Args:
        input_pdf_path: Path to the input PDF file
        output_pdf_path: Path for the output searchable PDF
        task_id: Task ID for progress tracking
        dpi: Resolution for rendering pages (higher = better OCR but slower)
        min_confidence: Minimum OCR confidence score (0.0-1.0)

    Returns:
        Dict with total_pages and output_path

    Raises:
        RuntimeError: If PDF is encrypted, corrupted, or processing is cancelled
    """
    output_path = Path(output_pdf_path)
    src_doc = None
    new_doc = None

    try:
        # Open and validate source PDF
        try:
            src_doc = fitz.open(input_pdf_path)
        except Exception as e:
            raise RuntimeError(f"無法開啟 PDF 檔案: {e}")

        # Check for encryption
        if src_doc.is_encrypted:
            src_doc.close()
            raise RuntimeError("無法處理加密的 PDF 檔案")

        total_pages = len(src_doc)
        if total_pages == 0:
            src_doc.close()
            raise RuntimeError("PDF 檔案沒有頁面")

        zoom = dpi / PDF_POINTS_PER_INCH

        # Create a new document
        new_doc = fitz.open()

        # Create font once outside the loop
        font = fitz.Font("cjk")

        # Copy metadata from source document
        metadata = src_doc.metadata
        if metadata:
            new_doc.set_metadata(metadata)

        update_progress(task_id, 0, total_pages, 'processing', f'開始處理 {total_pages} 頁...')
        failed_pages = []

        # Process pages in batches for better GPU utilization
        for batch_start in range(0, total_pages, OCR_BATCH_SIZE):
            # Check for cancellation at batch start
            if is_cancelled(task_id):
                raise RuntimeError("已取消處理")

            batch_end = min(batch_start + OCR_BATCH_SIZE, total_pages)
            batch_size = batch_end - batch_start

            update_progress(task_id, batch_start + 1, total_pages, 'processing',
                           f'辨識第 {batch_start + 1}-{batch_end}/{total_pages} 頁...')

            # Phase 1: Render all pages in batch to images
            # Note: We don't store pixmaps here to reduce memory pressure.
            # Pages are re-rendered in Phase 3 when needed for PDF insertion.
            batch_images = []
            batch_page_info = []  # Store page metadata for text overlay

            for page_idx in range(batch_start, batch_end):
                src_page = src_doc[page_idx]
                page_rect = src_page.rect
                rotation = src_page.rotation

                # Adaptive DPI: reduce for very large pages to save memory
                max_dimension = max(page_rect.width, page_rect.height)
                effective_dpi = dpi
                if max_dimension * zoom > 4000:
                    effective_dpi = int(4000 / max_dimension * 72)
                    effective_dpi = max(effective_dpi, 150)

                effective_zoom = effective_dpi / PDF_POINTS_PER_INCH

                # Render page to image for OCR
                matrix = fitz.Matrix(effective_zoom, effective_zoom)
                if rotation:
                    matrix = matrix.prerotate(rotation)

                pix = src_page.get_pixmap(matrix=matrix)
                img_array = pixmap_to_numpy(pix)

                batch_images.append(img_array)
                batch_page_info.append({
                    'page_idx': page_idx,
                    'page_rect': page_rect,
                    'rotation': rotation,
                    'effective_dpi': effective_dpi,
                    'effective_zoom': effective_zoom,
                })

                # Release pixmap immediately to reduce memory pressure
                del pix

            # Phase 2: Run batch OCR on all images
            try:
                batch_ocr_results = ocr_batch_images(batch_images, min_confidence=min_confidence)
            except Exception as e:
                logger.warning(f"OCR: Batch OCR failed, falling back to individual: {e}")
                batch_ocr_results = [ocr_image(img, min_confidence) for img in batch_images]
            finally:
                # Free batch images memory (always reached even if fallback fails)
                del batch_images

            # Phase 3: Create PDF pages with text overlay
            # Re-render pages here to avoid storing pixmaps in memory during OCR
            for i, (page_info, ocr_data) in enumerate(zip(batch_page_info, batch_ocr_results)):
                page_idx = page_info['page_idx']
                page_rect = page_info['page_rect']
                rotation = page_info['rotation']
                effective_dpi = page_info['effective_dpi']
                effective_zoom = page_info['effective_zoom']

                # Re-render page for PDF insertion (trades CPU for memory)
                src_page = src_doc[page_idx]
                matrix = fitz.Matrix(effective_zoom, effective_zoom)
                if rotation:
                    matrix = matrix.prerotate(rotation)
                pix = src_page.get_pixmap(matrix=matrix)

                try:
                    # Create new page with same dimensions
                    new_page = new_doc.new_page(width=page_rect.width, height=page_rect.height)

                    # Insert the original page as image (preserves visual appearance)
                    new_page.insert_image(page_rect, pixmap=pix)

                    # Calculate scale factors using effective zoom
                    img_width = page_rect.width * effective_zoom
                    img_height = page_rect.height * effective_zoom
                    pdf_width = page_rect.width
                    pdf_height = page_rect.height
                    scale_x = pdf_width / img_width
                    scale_y = pdf_height / img_height

                    # Row grouping tolerance scales with effective DPI (base 20px at 200 DPI)
                    row_tolerance = max(10, int(20 * (effective_dpi / 200)))

                    # Sort OCR results by position: top-to-bottom, then left-to-right
                    def get_sort_key(item):
                        poly = item['poly']
                        if len(poly) >= 4:
                            y_min = min(p[1] for p in poly)
                            x_min = min(p[0] for p in poly)
                            row = int(y_min / row_tolerance)
                            return (row, x_min)
                        return (0, 0)

                    ocr_data_sorted = sorted(ocr_data, key=get_sort_key)
                    text_errors = 0

                    for item in ocr_data_sorted:
                        text = item['text']
                        poly = item['poly']

                        if not text or len(poly) < 4:
                            continue

                        x_coords = [p[0] for p in poly]
                        y_coords = [p[1] for p in poly]

                        # Convert image coordinates to PDF coordinates
                        x0 = min(x_coords) * scale_x
                        y0 = min(y_coords) * scale_y
                        x1 = max(x_coords) * scale_x
                        y1 = max(y_coords) * scale_y

                        box_width = x1 - x0
                        box_height = y1 - y0

                        # Skip extremely small boxes
                        if box_width < 1 or box_height < 1:
                            continue

                        # Calculate font size to match text width to box width
                        ref_fontsize = 10.0
                        text_length = font.text_length(text, fontsize=ref_fontsize)
                        if text_length > 0:
                            fontsize = ref_fontsize * (box_width / text_length)
                        else:
                            fontsize = box_height * 0.8

                        # Clamp font size to reasonable range
                        fontsize = max(2.0, min(fontsize, 200.0))

                        try:
                            # Write each text block separately using its own TextWriter
                            tw = fitz.TextWriter(new_page.rect)
                            baseline_y = y1 - box_height * 0.15
                            tw.append((x0, baseline_y), text, font=font, fontsize=fontsize)
                            tw.write_text(new_page, render_mode=3)  # Invisible text
                        except Exception as e:
                            text_errors += 1
                            if text_errors <= 3:
                                logger.warning(f"OCR: Failed to write text on page {page_idx + 1}: {e}")

                    if text_errors > 3:
                        logger.warning(f"OCR: {text_errors} text insertion errors on page {page_idx + 1}")

                except Exception as e:
                    logger.error(f"OCR: Error processing page {page_idx + 1}: {e}")
                    failed_pages.append(page_idx + 1)
                    # Still create the page with just the image (no text layer)
                    try:
                        if new_doc.page_count <= page_idx:
                            new_page = new_doc.new_page(width=page_rect.width, height=page_rect.height)
                            new_page.insert_image(page_rect, pixmap=pix)
                    except Exception:
                        pass
                finally:
                    # Free pixmap memory after use
                    del pix

            del batch_page_info

        src_doc.close()
        src_doc = None

        update_progress(task_id, total_pages, total_pages, 'saving', '儲存 PDF 中...')
        new_doc.save(output_pdf_path, garbage=4, deflate=True)
        new_doc.close()
        new_doc = None

        update_progress(task_id, total_pages, total_pages, 'done', '完成!')

        result = {
            'total_pages': total_pages,
            'output_path': output_pdf_path
        }

        if failed_pages:
            result['warning'] = f"部分頁面處理失敗: {failed_pages}"
            logger.warning(f"OCR: Completed with warnings. Failed pages: {failed_pages}")

        return result

    except Exception as e:
        # Clean up on any error
        if src_doc:
            try:
                src_doc.close()
            except Exception:
                pass

        if new_doc:
            try:
                new_doc.close()
            except Exception:
                pass

        # Remove partial output file
        if output_path.exists():
            try:
                output_path.unlink()
                logger.info(f"OCR: Cleaned up partial output: {output_path}")
            except Exception:
                pass

        raise


def convert_pdf_to_markdown(input_pdf_path: str, output_dir: Path, task_id: str, original_filename: str = None) -> dict:
    """Convert PDF to Markdown using PaddleOCR-VL-1.5."""

    if original_filename:
        base_name = Path(original_filename).stem
    else:
        base_name = Path(input_pdf_path).stem

    # Use shared VL processing (now returns download_id for unique folder naming)
    restructured, file_output_dir, total_pages, download_id = process_pdf_with_vl(
        input_pdf_path, output_dir, task_id, original_filename
    )

    try:
        # Save to markdown
        for res in restructured:
            res.save_to_markdown(save_path=str(file_output_dir))

        # Find and concatenate all markdown files (VL may produce multiple)
        md_files = sorted(file_output_dir.glob("*.md"))
        if not md_files:
            raise RuntimeError("VL 模型未產生 Markdown 檔案")

        # Concatenate all markdown files if multiple exist
        markdown_parts = []
        for md_file in md_files:
            with open(md_file, "r", encoding="utf-8") as f:
                markdown_parts.append(f.read())
        markdown_text = "\n\n".join(markdown_parts)

        # Fix simplified Chinese characters
        markdown_text = fix_ocr_text(markdown_text)

        # Remove generic "Image" alt text from VL model HTML output (provides no useful info)
        markdown_text = _re.sub(r'alt="Image"', 'alt=""', markdown_text, flags=_re.IGNORECASE)

        # Convert HTML tables to markdown tables (cleaner output, filters empty rows)
        for _ in range(100):  # Safety limit
            found_match = False
            for pattern in _TABLE_PATTERNS:
                match = pattern.search(markdown_text)
                if match:
                    html_chunk = match.group(0)
                    md_table = html_table_to_markdown(html_chunk)
                    markdown_text = markdown_text[:match.start()] + md_table + markdown_text[match.end():]
                    found_match = True
                    break
            if not found_match:
                break

        # Save to final path (use download_id for consistency)
        final_md_path = file_output_dir / f"{download_id}.md"
        with open(final_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)

        # Clean up original md files
        for md_file in md_files:
            if md_file != final_md_path and md_file.exists():
                md_file.unlink()

        # Count images in output directory (include all common formats)
        image_patterns = ["**/*.png", "**/*.jpg", "**/*.jpeg", "**/*.gif", "**/*.webp"]
        saved_images = []
        for pattern in image_patterns:
            saved_images.extend(file_output_dir.glob(pattern))

        update_progress(task_id, total_pages, total_pages, 'done', '完成!')

        return {
            'total_pages': total_pages,
            'markdown_path': str(final_md_path),
            'output_dir': str(file_output_dir),
            'images': [str(img.relative_to(file_output_dir)) for img in saved_images],
            'download_id': download_id
        }

    except Exception:
        # Clean up on failure
        if file_output_dir.exists():
            shutil.rmtree(file_output_dir)
        raise


def _fix_word_styles(docx_path: Path):
    """Fix Word document styles - remove blue color and add table borders.

    Uses atomic write pattern: saves to temp file first, then replaces original.
    This prevents corruption if the save operation fails mid-write.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document(str(docx_path))

        # Fix paragraph styles (headings and body text)
        for para in doc.paragraphs:
            for run in para.runs:
                try:
                    # Force all text to black (fixes pandoc's blue headings)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                except Exception:
                    pass  # Skip malformed runs

        # Fix table cell styles and add borders
        for table in doc.tables:
            try:
                # Add table borders
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                if tbl.tblPr is None:
                    tbl.insert(0, tblPr)

                # Create table borders element
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # border width
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')  # black
                    tblBorders.append(border)

                # Remove existing borders and add new ones
                existing_borders = tblPr.find(qn('w:tblBorders'))
                if existing_borders is not None:
                    tblPr.remove(existing_borders)
                tblPr.append(tblBorders)

                # Fix cell text colors
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                try:
                                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                                except Exception:
                                    pass
            except Exception as e:
                logger.warning(f"Word: Could not fix table borders: {e}")
                continue

        # Atomic save: write to temp file first, then replace original
        temp_path = docx_path.with_suffix('.docx.tmp')
        try:
            doc.save(str(temp_path))
            # Atomic replace (on POSIX systems, rename is atomic)
            temp_path.replace(docx_path)
        except Exception:
            # Clean up temp file on failure
            if temp_path.exists():
                temp_path.unlink()
            raise

    except Exception as e:
        # Don't fail the entire conversion if style fixing fails
        logger.warning(f"Word: Could not fix Word styles: {e}")


def html_table_to_markdown(html_content: str) -> str:
    """Convert HTML table to Markdown table using BeautifulSoup.

    Handles colspan by using empty cells (not duplicating content).
    Preserves all rows including empty ones to maintain original format.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html_content, 'html.parser')
    table = soup.find('table')
    if not table:
        return html_content

    rows = table.find_all('tr')
    if not rows:
        return html_content

    md_rows = []
    max_cols = 0

    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_cells = []

        for cell in cells:
            # Handle colspan - use empty cells for spanned columns instead of duplicating
            colspan = int(cell.get('colspan', 1))

            # Convert <br> to space, get text content
            for br in cell.find_all('br'):
                br.replace_with(' ')

            # Get text and clean it
            text = cell.get_text(separator=' ', strip=True)
            text = text.replace('|', '\\|').replace('\n', ' ')
            text = ' '.join(text.split())  # Normalize whitespace

            # Add cell content, then empty cells for remaining colspan
            row_cells.append(text if text else ' ')
            for _ in range(colspan - 1):
                row_cells.append(' ')

        # Keep all rows to preserve original format
        if row_cells:
            md_rows.append(row_cells)
            max_cols = max(max_cols, len(row_cells))

    if not md_rows:
        return html_content

    # Normalize all rows to same column count
    for row in md_rows:
        while len(row) < max_cols:
            row.append(' ')

    # Build markdown table
    result = []
    result.append('| ' + ' | '.join(md_rows[0]) + ' |')
    result.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in md_rows[1:]:
        result.append('| ' + ' | '.join(row) + ' |')

    return '\n\n' + '\n'.join(result) + '\n\n'


def html_img_to_markdown(html_content: str, output_dir: Path = None) -> str:
    """Convert HTML img tag to Markdown image syntax, preserving alt text.

    Args:
        html_content: HTML string containing an img tag
        output_dir: Optional directory to save extracted data URI images

    Returns:
        Markdown image syntax with properly escaped alt text and URL
    """
    from bs4 import BeautifulSoup
    import base64
    import re
    from urllib.parse import quote

    soup = BeautifulSoup(html_content, 'html.parser')
    img = soup.find('img')
    if not img:
        return html_content

    src = img.get('src', '')
    alt = img.get('alt', '')

    if not src:
        return html_content

    # Clear generic "Image" alt text from VL model (provides no useful information)
    if alt.strip().lower() == 'image':
        alt = ''

    # Escape special markdown characters in alt text: [ ] \
    alt = alt.replace('\\', '\\\\').replace('[', '\\[').replace(']', '\\]')

    # Handle data URIs - extract to file if output_dir provided, otherwise use placeholder
    if src.startswith('data:'):
        if output_dir and output_dir.is_dir():
            # Extract data URI to file
            # Format: data:[<mediatype>][;base64],<data>
            match = re.match(r'data:([^;,]+)?(?:;base64)?,(.+)', src, re.DOTALL)
            if match:
                media_type = match.group(1) or 'application/octet-stream'
                data = match.group(2)

                # Determine file extension from media type
                ext_map = {
                    'image/png': '.png',
                    'image/jpeg': '.jpg',
                    'image/gif': '.gif',
                    'image/webp': '.webp',
                    'image/svg+xml': '.svg',
                }
                ext = ext_map.get(media_type, '.bin')

                # Generate unique filename
                img_id = str(uuid.uuid4())[:8]
                img_filename = f"extracted_{img_id}{ext}"
                img_path = output_dir / img_filename

                try:
                    # Decode and save
                    img_data = base64.b64decode(data)
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    src = img_filename
                except Exception:
                    # Fall back to placeholder if extraction fails
                    return f'\n\n<!-- Image: {alt or "embedded image"} (data URI not extracted) -->\n\n'
            else:
                return f'\n\n<!-- Image: {alt or "embedded image"} (invalid data URI) -->\n\n'
        else:
            # No output dir - use placeholder comment instead of huge inline data
            return f'\n\n<!-- Image: {alt or "embedded image"} (embedded data URI) -->\n\n'

    # Escape special characters in URL: spaces and parentheses
    # Parentheses break markdown link syntax, spaces need encoding
    src = src.replace(' ', '%20').replace('(', '%28').replace(')', '%29')

    return f'\n\n![{alt}]({src})\n\n'


def process_markdown_for_word(markdown_text: str, output_dir: Path = None) -> str:
    """Process markdown text to convert HTML elements for Word compatibility.

    Args:
        markdown_text: Raw markdown text potentially containing HTML elements
        output_dir: Optional directory for extracting data URI images

    Returns:
        Processed markdown with HTML tables/images converted to markdown syntax
    """
    # Fix simplified Chinese characters
    markdown_text = fix_ocr_text(markdown_text)

    # Convert HTML tables - loop until no more matches to handle nested cases
    # (e.g., a table replacement might reveal another table pattern)
    max_iterations = 100  # Safety limit to prevent infinite loops
    for _ in range(max_iterations):
        found_match = False
        for pattern in _TABLE_PATTERNS:
            match = pattern.search(markdown_text)
            if match:
                html_chunk = match.group(0)
                md_table = html_table_to_markdown(html_chunk)
                markdown_text = markdown_text[:match.start()] + md_table + markdown_text[match.end():]
                found_match = True
                break  # Re-scan from beginning after each replacement
        if not found_match:
            break

    # Convert HTML img tags - loop until no more matches
    # This handles cases where table conversion might have revealed image tags
    for _ in range(max_iterations):
        found_match = False
        for pattern in _IMG_PATTERNS:
            match = pattern.search(markdown_text)
            if match:
                html_chunk = match.group(0)
                md_img = html_img_to_markdown(html_chunk, output_dir)
                markdown_text = markdown_text[:match.start()] + md_img + markdown_text[match.end():]
                found_match = True
                break  # Re-scan from beginning after each replacement
        if not found_match:
            break

    return markdown_text


def convert_pdf_to_word(input_pdf_path: str, output_dir: Path, task_id: str, original_filename: str = None) -> dict:
    """Convert PDF to Word document using PaddleOCR-VL-1.5 + pandoc."""
    import subprocess

    if original_filename:
        base_name = Path(original_filename).stem
    else:
        base_name = Path(input_pdf_path).stem

    # Word conversion has 2 extra steps after VL processing:
    # 1. Processing tables/images
    # 2. Pandoc conversion
    extra_steps = 2

    # Use shared VL processing with extra_steps for consistent progress tracking
    restructured, file_output_dir, total_pages, download_id = process_pdf_with_vl(
        input_pdf_path, output_dir, task_id, original_filename, extra_steps=extra_steps
    )

    # Total steps = VL pages + post-processing steps
    total_steps = total_pages + extra_steps

    try:
        # Save to markdown
        for res in restructured:
            res.save_to_markdown(save_path=str(file_output_dir))

        # Find and concatenate all markdown files (VL may produce multiple)
        md_files = sorted(file_output_dir.glob("*.md"))
        if not md_files:
            raise RuntimeError("VL 模型未產生 Markdown 檔案")

        # Concatenate all markdown files if multiple exist
        markdown_parts = []
        for md_file in md_files:
            with open(md_file, "r", encoding="utf-8") as f:
                markdown_parts.append(f.read())
        markdown_text = "\n\n".join(markdown_parts)

        # Process HTML elements in markdown
        update_progress(task_id, total_pages + 1, total_steps, 'processing', '處理表格與圖片...')
        markdown_text = process_markdown_for_word(markdown_text, file_output_dir)

        # Save processed markdown with correct name (use download_id for consistency)
        final_md_path = file_output_dir / f"{download_id}.md"
        with open(final_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        # Clean up original md files
        for md_file in md_files:
            if md_file != final_md_path and md_file.exists():
                md_file.unlink()

        update_progress(task_id, total_pages + 1, total_steps, 'converting', '轉換為 Word 中...')

        # Convert Markdown to Word using pandoc
        docx_filename = f"{download_id}.docx"
        docx_path = file_output_dir / docx_filename

        try:
            result = subprocess.run(
                ['pandoc', str(final_md_path), '-o', str(docx_path),
                 '--resource-path', str(file_output_dir),
                 '--extract-media', str(file_output_dir)],
                check=True,
                capture_output=True,
                cwd=str(file_output_dir),
                timeout=PANDOC_TIMEOUT
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"Pandoc conversion timed out after {PANDOC_TIMEOUT} seconds")
        except subprocess.CalledProcessError as e:
            error_detail = f"stdout: {e.stdout.decode()}, stderr: {e.stderr.decode()}"
            raise RuntimeError(f"Pandoc conversion failed: {error_detail}")
        except FileNotFoundError:
            raise RuntimeError("Pandoc not installed. Please install pandoc.")

        # Post-process: fix blue headings and ensure proper styling
        _fix_word_styles(docx_path)

        # Count images in output directory (include all common formats)
        image_patterns = ["**/*.png", "**/*.jpg", "**/*.jpeg", "**/*.gif", "**/*.webp"]
        images = []
        for pattern in image_patterns:
            images.extend(file_output_dir.glob(pattern))

        update_progress(task_id, total_steps, total_steps, 'done', '完成!')

        return {
            'total_pages': total_pages,
            'output_path': str(docx_path),
            'download_id': download_id,
            'images_count': len(images)
        }

    except Exception:
        # Clean up on failure
        if file_output_dir.exists():
            shutil.rmtree(file_output_dir)
        raise


# ============== Background Processing Functions ==============

def process_ocr_job(job_id: str, input_path: str, original_filename: str):
    """Background worker for OCR processing."""
    # Wait for OCR lock (allows VL jobs to run in parallel)
    update_progress(job_id, 0, 1, 'waiting', '等待其他任務完成...')

    output_path = None

    with ocr_processing_lock:
        try:
            # Check if cancelled while waiting
            if is_cancelled(job_id):
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
                return

            file_id = str(uuid.uuid4())[:8]
            base_name = Path(original_filename).stem
            output_filename = f"{base_name}_searchable_{file_id}.pdf"
            output_path = OUTPUT_DIR / output_filename

            result = create_searchable_pdf(input_path, str(output_path), job_id)

            job_result = {
                'total_pages': result['total_pages'],
                'download_id': output_filename,
                'original_name': original_filename
            }

            # Include warning if some pages failed
            if 'warning' in result:
                job_result['warning'] = result['warning']

            update_job(job_id, status='done', result=job_result)

        except Exception as e:
            import traceback
            error_msg = str(e)

            # Clean up output file on error (create_searchable_pdf should do this,
            # but double-check in case of unexpected errors)
            if output_path and output_path.exists():
                try:
                    output_path.unlink()
                except Exception:
                    pass

            if '已取消' in error_msg:
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
            else:
                update_progress(job_id, 0, 0, 'error', error_msg)
                update_job(job_id,
                    status='error',
                    error=error_msg,
                    traceback=traceback.format_exc()
                )
        finally:
            # Clean up cancel flag
            with cancel_lock:
                cancel_flags.pop(job_id, None)
            # Clean up uploaded file
            try:
                os.unlink(input_path)
            except Exception as e:
                logger.warning(f"Cleanup: Failed to delete upload {input_path}: {e}")


def process_markdown_job(job_id: str, input_path: str, original_filename: str):
    """Background worker for Markdown conversion."""
    # Wait for VL lock (allows OCR jobs to run in parallel)
    update_progress(job_id, 0, 1, 'waiting', '等待其他任務完成...')

    with vl_processing_lock:
        try:
            # Check if cancelled while waiting
            if is_cancelled(job_id):
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
                return

            result = convert_pdf_to_markdown(input_path, OUTPUT_DIR, job_id, original_filename)

            update_job(job_id,
                status='done',
                result={
                    'total_pages': result['total_pages'],
                    'download_id': result['download_id'],
                    'images_count': len(result['images']),
                    'original_name': original_filename
                }
            )
        except Exception as e:
            import traceback
            error_msg = str(e)
            if '已取消' in error_msg:
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
            else:
                update_progress(job_id, 0, 0, 'error', error_msg)
                update_job(job_id,
                    status='error',
                    error=error_msg,
                    traceback=traceback.format_exc()
                )
        finally:
            # Clean up cancel flag
            with cancel_lock:
                cancel_flags.pop(job_id, None)
            try:
                os.unlink(input_path)
            except Exception as e:
                logger.warning(f"Cleanup: Failed to delete upload {input_path}: {e}")


def process_word_job(job_id: str, input_path: str, original_filename: str):
    """Background worker for Word conversion."""
    # Wait for VL lock (allows OCR jobs to run in parallel)
    update_progress(job_id, 0, 1, 'waiting', '等待其他任務完成...')

    with vl_processing_lock:
        try:
            # Check if cancelled while waiting
            if is_cancelled(job_id):
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
                return

            result = convert_pdf_to_word(input_path, OUTPUT_DIR, job_id, original_filename)

            update_job(job_id,
                status='done',
                result={
                    'total_pages': result['total_pages'],
                    'download_id': result['download_id'],
                    'images_count': result.get('images_count', 0),
                    'original_name': original_filename
                }
            )
        except Exception as e:
            import traceback
            error_msg = str(e)
            if '已取消' in error_msg:
                update_progress(job_id, 0, 0, 'cancelled', '已取消')
                update_job(job_id, status='cancelled')
            else:
                update_progress(job_id, 0, 0, 'error', error_msg)
                update_job(job_id,
                    status='error',
                    error=error_msg,
                    traceback=traceback.format_exc()
                )
        finally:
            # Clean up cancel flag
            with cancel_lock:
                cancel_flags.pop(job_id, None)
            try:
                os.unlink(input_path)
            except Exception as e:
                logger.warning(f"Cleanup: Failed to delete upload {input_path}: {e}")


def process_dual_export_job(md_job_id: str, word_job_id: str, input_path: str, original_filename: str):
    """Background worker for dual Markdown+Word export. VL runs once, both jobs updated."""
    import subprocess

    # Both jobs show waiting initially
    update_progress(md_job_id, 0, 1, 'waiting', '等待其他任務完成...')
    update_progress(word_job_id, 0, 1, 'waiting', '等待其他任務完成...')

    with vl_processing_lock:
        try:
            # Check if either job cancelled
            if is_cancelled(md_job_id) or is_cancelled(word_job_id):
                for jid in [md_job_id, word_job_id]:
                    update_progress(jid, 0, 0, 'cancelled', '已取消')
                    update_job(jid, status='cancelled')
                return

            # Phase 1: VL Processing (runs once for both outputs)
            # Word has 2 extra steps after VL; markdown completes after VL
            # Use extra_steps=2 so Word job progress is consistent
            extra_steps = 2
            restructured, file_output_dir, total_pages, download_id = process_pdf_with_vl(
                input_path, OUTPUT_DIR, md_job_id, original_filename, extra_steps=extra_steps
            )

            total_steps = total_pages + extra_steps

            # Sync word job progress with markdown job
            with progress_lock:
                if md_job_id in progress_data:
                    progress_data[word_job_id] = progress_data[md_job_id].copy()

            # Phase 2: Save Markdown
            # Markdown job uses total_pages as its total (no extra steps)
            update_progress(md_job_id, total_pages, total_pages, 'saving', '儲存 Markdown...')
            # Word job uses total_steps (with extra steps)
            update_progress(word_job_id, total_pages, total_steps, 'saving', '儲存 Markdown...')

            for res in restructured:
                res.save_to_markdown(save_path=str(file_output_dir))

            # Find and concatenate all markdown files (VL may produce multiple)
            md_files = sorted(file_output_dir.glob("*.md"))
            if not md_files:
                raise RuntimeError("VL 模型未產生 Markdown 檔案")

            # Concatenate all markdown files if multiple exist
            markdown_parts = []
            for md_file in md_files:
                with open(md_file, "r", encoding="utf-8") as f:
                    markdown_parts.append(f.read())
            markdown_text = "\n\n".join(markdown_parts)

            markdown_text = fix_ocr_text(markdown_text)

            # Remove generic "Image" alt text from VL model HTML output
            markdown_text = _re.sub(r'alt="Image"', 'alt=""', markdown_text, flags=_re.IGNORECASE)

            final_md_path = file_output_dir / f"{download_id}.md"
            with open(final_md_path, "w", encoding="utf-8") as f:
                f.write(markdown_text)

            # Clean up original md files
            for md_file in md_files:
                if md_file != final_md_path and md_file.exists():
                    md_file.unlink()

            # Count images in output directory
            image_patterns = ["**/*.png", "**/*.jpg", "**/*.jpeg", "**/*.gif", "**/*.webp"]
            images = []
            for pattern in image_patterns:
                images.extend(file_output_dir.glob(pattern))

            # Markdown job done
            update_progress(md_job_id, total_pages, total_pages, 'done', '完成!')
            update_job(md_job_id, status='done', result={
                'total_pages': total_pages,
                'download_id': download_id,
                'images_count': len(images),
                'original_name': original_filename
            })

            # Phase 3: Convert to Word
            update_progress(word_job_id, total_pages + 1, total_steps, 'converting', '轉換為 Word...')

            # Process markdown for Word (HTML tables -> MD tables)
            with open(final_md_path, "r", encoding="utf-8") as f:
                word_markdown = f.read()
            word_markdown = process_markdown_for_word(word_markdown, file_output_dir)

            # Save processed markdown for pandoc
            word_md_path = file_output_dir / f"{download_id}_word.md"
            with open(word_md_path, "w", encoding="utf-8") as f:
                f.write(word_markdown)

            # Pandoc conversion
            docx_path = file_output_dir / f"{download_id}.docx"
            try:
                subprocess.run(
                    ['pandoc', str(word_md_path), '-o', str(docx_path),
                     '--resource-path', str(file_output_dir),
                     '--extract-media', str(file_output_dir)],
                    check=True, capture_output=True,
                    cwd=str(file_output_dir), timeout=PANDOC_TIMEOUT
                )
            except subprocess.TimeoutExpired:
                raise RuntimeError(f"Pandoc conversion timed out after {PANDOC_TIMEOUT} seconds")
            except subprocess.CalledProcessError as e:
                error_detail = f"stdout: {e.stdout.decode()}, stderr: {e.stderr.decode()}"
                raise RuntimeError(f"Pandoc conversion failed: {error_detail}")
            except FileNotFoundError:
                raise RuntimeError("Pandoc not installed. Please install pandoc.")

            _fix_word_styles(docx_path)

            # Clean up temp markdown
            if word_md_path.exists():
                word_md_path.unlink()

            # Word job done
            update_progress(word_job_id, total_steps, total_steps, 'done', '完成!')
            update_job(word_job_id, status='done', result={
                'total_pages': total_pages,
                'download_id': download_id,
                'images_count': len(images),
                'original_name': original_filename
            })

        except Exception as e:
            import traceback
            error_msg = str(e)
            if '已取消' in error_msg:
                for jid in [md_job_id, word_job_id]:
                    update_progress(jid, 0, 0, 'cancelled', '已取消')
                    update_job(jid, status='cancelled')
            else:
                for jid in [md_job_id, word_job_id]:
                    update_progress(jid, 0, 0, 'error', error_msg)
                    update_job(jid, status='error', error=error_msg,
                              traceback=traceback.format_exc())
        finally:
            for jid in [md_job_id, word_job_id]:
                with cancel_lock:
                    cancel_flags.pop(jid, None)
            try:
                os.unlink(input_path)
            except Exception as e:
                logger.warning(f"Cleanup: Failed to delete upload {input_path}: {e}")


# ============== Routes ==============

@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


@app.route('/api/progress/<task_id>')
def get_progress(task_id):
    """SSE endpoint for progress updates."""
    def generate():
        last_data = None
        start_time = time.time()
        max_duration = 3600  # 1 hour timeout to prevent infinite loops

        while time.time() - start_time < max_duration:
            with progress_lock:
                data = progress_data.get(task_id, {
                    'current': 0, 'total': 0, 'percent': 0,
                    'status': 'waiting', 'message': '等待中...'
                })

            # Only send if data changed
            if data != last_data:
                yield f"data: {json.dumps(data)}\n\n"
                last_data = data.copy()

                if data.get('status') in ('done', 'error', 'cancelled'):
                    break

            time.sleep(0.3)
        else:
            # Timeout reached - notify client
            yield f"data: {json.dumps({'status': 'timeout', 'message': '連線逾時'})}\n\n"

    return Response(generate(), mimetype='text/event-stream',
                   headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/job/<job_id>')
def get_job_status(job_id):
    """Get job status and result."""
    with jobs_lock:
        job = jobs.get(job_id)

    if not job:
        return jsonify({'error': 'Job not found'}), 404

    # Include progress data
    with progress_lock:
        progress = progress_data.get(job_id, {})

    response = jsonify({
        'job_id': job_id,
        'status': job['status'],
        'filename': job['filename'],
        'mode': job['mode'],
        'progress': progress,
        'result': job.get('result'),
        'error': job.get('error'),
        'created_at': job['created_at'],
        'updated_at': job.get('updated_at')
    })
    # Prevent caching to ensure fresh progress data
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/api/ocr', methods=['POST'])
@limiter.limit("5 per minute")
def ocr_endpoint():
    """Accept PDF upload and start async OCR processing."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    job_ids = []

    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            continue

        # Validate PDF magic bytes
        if not validate_pdf_file(file):
            return jsonify({'error': f'File {file.filename} is not a valid PDF'}), 400

        # Generate job ID
        job_id = str(uuid.uuid4())[:8]

        # Save uploaded file
        upload_path = UPLOAD_DIR / f"{job_id}.pdf"
        file.save(str(upload_path))

        # Create job entry
        with jobs_lock:
            jobs[job_id] = {
                'status': 'processing',
                'filename': file.filename,
                'mode': 'ocr',
                'created_at': time.time(),
                'updated_at': time.time()
            }

        # Start background processing
        thread = threading.Thread(
            target=process_ocr_job,
            args=(job_id, str(upload_path), file.filename)
        )
        thread.daemon = True
        thread.start()

        job_ids.append({
            'job_id': job_id,
            'filename': file.filename
        })

    if not job_ids:
        return jsonify({'error': 'No valid PDF files provided'}), 400

    return jsonify({'jobs': job_ids})


@app.route('/api/markdown', methods=['POST'])
@limiter.limit("5 per minute")
def markdown_endpoint():
    """Accept PDF upload and start async Markdown conversion."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    job_ids = []

    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            continue

        # Validate PDF magic bytes
        if not validate_pdf_file(file):
            return jsonify({'error': f'File {file.filename} is not a valid PDF'}), 400

        job_id = str(uuid.uuid4())[:8]
        upload_path = UPLOAD_DIR / f"{job_id}.pdf"
        file.save(str(upload_path))

        with jobs_lock:
            jobs[job_id] = {
                'status': 'processing',
                'filename': file.filename,
                'mode': 'markdown',
                'created_at': time.time(),
                'updated_at': time.time()
            }

        thread = threading.Thread(
            target=process_markdown_job,
            args=(job_id, str(upload_path), file.filename)
        )
        thread.daemon = True
        thread.start()

        job_ids.append({
            'job_id': job_id,
            'filename': file.filename
        })

    if not job_ids:
        return jsonify({'error': 'No valid PDF files provided'}), 400

    return jsonify({'jobs': job_ids})


@app.route('/api/word', methods=['POST'])
@limiter.limit("5 per minute")
def word_endpoint():
    """Accept PDF upload and start async Word conversion."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    job_ids = []

    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            continue

        # Validate PDF magic bytes
        if not validate_pdf_file(file):
            return jsonify({'error': f'File {file.filename} is not a valid PDF'}), 400

        job_id = str(uuid.uuid4())[:8]
        upload_path = UPLOAD_DIR / f"{job_id}.pdf"
        file.save(str(upload_path))

        with jobs_lock:
            jobs[job_id] = {
                'status': 'processing',
                'filename': file.filename,
                'mode': 'word',
                'created_at': time.time(),
                'updated_at': time.time()
            }

        thread = threading.Thread(
            target=process_word_job,
            args=(job_id, str(upload_path), file.filename)
        )
        thread.daemon = True
        thread.start()

        job_ids.append({
            'job_id': job_id,
            'filename': file.filename
        })

    if not job_ids:
        return jsonify({'error': 'No valid PDF files provided'}), 400

    return jsonify({'jobs': job_ids})


@app.route('/api/export', methods=['POST'])
@limiter.limit("5 per minute")
def export_endpoint():
    """Dual export: Markdown + Word from single VL run. Returns two job IDs per file."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    job_pairs = []

    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            continue

        if not validate_pdf_file(file):
            return jsonify({'error': f'File {file.filename} is not a valid PDF'}), 400

        # Create two job IDs for same file
        md_job_id = str(uuid.uuid4())[:8]
        word_job_id = str(uuid.uuid4())[:8]

        upload_path = UPLOAD_DIR / f"{md_job_id}.pdf"
        file.save(str(upload_path))

        now = time.time()
        with jobs_lock:
            jobs[md_job_id] = {
                'status': 'processing',
                'filename': file.filename,
                'mode': 'markdown',
                'created_at': now,
                'updated_at': now
            }
            jobs[word_job_id] = {
                'status': 'processing',
                'filename': file.filename,
                'mode': 'word',
                'created_at': now,
                'updated_at': now
            }

        thread = threading.Thread(
            target=process_dual_export_job,
            args=(md_job_id, word_job_id, str(upload_path), file.filename)
        )
        thread.daemon = True
        thread.start()

        job_pairs.append({
            'filename': file.filename,
            'jobs': [
                {'job_id': md_job_id, 'mode': 'markdown'},
                {'job_id': word_job_id, 'mode': 'word'}
            ]
        })

    if not job_pairs:
        return jsonify({'error': 'No valid PDF files provided'}), 400

    return jsonify({'job_pairs': job_pairs})


@app.route('/api/download/word/<folder_name>')
def download_word(folder_name):
    """Download Word document."""
    try:
        folder_path = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    if not folder_path.exists():
        return jsonify({'error': 'File not found'}), 404

    # Check that it's actually a directory
    if not folder_path.is_dir():
        return jsonify({'error': 'Not a directory'}), 400

    docx_path = folder_path / f"{folder_name}.docx"

    if not docx_path.exists():
        return jsonify({'error': 'Word file not found'}), 404

    # No auto-delete - user can dismiss manually or wait for scheduled cleanup
    # Use original filename if provided via query param, otherwise use folder name
    original_name = request.args.get('name', folder_name)
    # Extract stem (remove extension) and sanitize for HTTP header safety
    base_name = Path(original_name).stem
    safe_name = sanitize_download_name(base_name)
    return send_file(
        docx_path,
        as_attachment=True,
        download_name=f"{safe_name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/download/markdown/<folder_name>')
def download_markdown(folder_name):
    """Download markdown folder as zip."""
    import zipfile

    try:
        folder_path = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    if not folder_path.exists():
        return jsonify({'error': 'File not found'}), 404

    # Check that it's actually a directory
    if not folder_path.is_dir():
        return jsonify({'error': 'Not a directory'}), 400

    # Create zip file with unique name to avoid race conditions
    # Use uuid suffix to prevent concurrent downloads from overwriting each other
    zip_id = str(uuid.uuid4())[:8]
    zip_filename = f"{folder_name}_{zip_id}.zip"
    zip_path = OUTPUT_DIR / zip_filename

    # User-friendly download name (without the uuid suffix)
    # Use original filename if provided via query param, otherwise use folder name
    original_name = request.args.get('name', folder_name)
    # Extract stem (remove extension) and sanitize for HTTP header safety
    base_name = Path(original_name).stem
    safe_name = sanitize_download_name(base_name)
    download_name = f"{safe_name}.zip"

    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in folder_path.rglob('*'):
                # Exclude Word files from markdown download
                if file_path.is_file() and not file_path.suffix.lower() == '.docx':
                    arcname = file_path.relative_to(folder_path)
                    zipf.write(file_path, arcname)

        # Schedule zip cleanup after download (zip is temporary, folder persists)
        @after_this_request
        def cleanup_zip(response):
            schedule_file_deletion(zip_path, delay=10.0)
            return response

        return send_file(
            zip_path,
            as_attachment=True,
            download_name=download_name,
            mimetype='application/zip'
        )
    except Exception:
        # Clean up zip file if an error occurs before response is sent
        if zip_path.exists():
            try:
                zip_path.unlink()
            except Exception:
                pass
        raise


@app.route('/api/view/markdown/<folder_name>')
def view_markdown(folder_name):
    """View markdown content."""
    try:
        folder_path = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    if not folder_path.exists():
        return jsonify({'error': 'File not found'}), 404

    # Check that it's actually a directory
    if not folder_path.is_dir():
        return jsonify({'error': 'Not a directory'}), 400

    md_path = folder_path / f"{folder_name}.md"

    if not md_path.exists():
        return jsonify({'error': 'Markdown file not found'}), 404

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    return jsonify({
        'content': content,
        'filename': f"{folder_name}.md"
    })


@app.route('/api/download/<filename>')
def download_file(filename):
    """Download the processed PDF file."""
    try:
        file_path = validate_path(filename)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    if not file_path.exists():
        return jsonify({'error': 'File not found'}), 404

    # Use original filename if provided, otherwise use the stored filename
    # Sanitize for HTTP header safety (user-provided query param)
    original_name = request.args.get('name', filename)
    safe_name = sanitize_download_name(original_name)
    # Ensure .pdf extension
    if not safe_name.lower().endswith('.pdf'):
        safe_name = f"{safe_name}.pdf"

    # No auto-delete - user can dismiss manually or wait for scheduled cleanup
    return send_file(
        file_path,
        as_attachment=True,
        download_name=safe_name,
        mimetype='application/pdf'
    )


@app.route('/api/view/<filename>')
def view_file(filename):
    """View the processed PDF file inline."""
    try:
        file_path = validate_path(filename)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    if not file_path.exists():
        return jsonify({'error': 'File not found'}), 404

    return send_file(
        file_path,
        as_attachment=False,
        mimetype='application/pdf'
    )


@app.route('/api/delete/<path:file_id>', methods=['DELETE'])
def delete_export(file_id):
    """Delete exported file or folder."""
    try:
        resolved_path = validate_path(file_id)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400

    try:
        # Try as PDF file
        if resolved_path.exists() and resolved_path.is_file():
            resolved_path.unlink()
            logger.info(f"Cleanup: User deleted: {file_id}")
            return jsonify({'status': 'deleted', 'file': file_id})

        # Try as folder (markdown/word)
        if resolved_path.exists() and resolved_path.is_dir():
            shutil.rmtree(resolved_path)
            logger.info(f"Cleanup: User deleted folder: {file_id}/")
            # Also delete zip if exists
            zip_path = OUTPUT_DIR / f"{file_id}.zip"
            if zip_path.exists():
                zip_path.unlink()
            return jsonify({'status': 'deleted', 'folder': file_id})

        return jsonify({'status': 'not_found'}), 404
    except Exception as e:
        logger.info(f"Cleanup: Error deleting {file_id}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/cancel/<job_id>', methods=['POST'])
def cancel_job(job_id):
    """Cancel a processing job."""
    # Keep entire operation within lock to prevent race conditions
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404
        if job['status'] != 'processing':
            return jsonify({'error': 'Job is not processing'}), 400
        # Set directly to 'cancelled' to match progress status
        job['status'] = 'cancelled'
        job['updated_at'] = time.time()

    # Set cancel flag and update progress outside lock
    set_cancelled(job_id, True)
    update_progress(job_id, 0, 0, 'cancelled', '已取消')

    return jsonify({'status': 'cancelled', 'job_id': job_id})


@app.route('/api/health')
def health():
    return jsonify({'status': 'ok'})


@app.route('/api/csrf-token')
def get_csrf_token():
    """Get CSRF token for SPA."""
    return jsonify({'csrf_token': generate_csrf()})


if __name__ == '__main__':
    # Development only - production uses gunicorn
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)
